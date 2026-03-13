"""Word (.docx) formatter for the Governance Agent output."""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


def _set_heading_color(paragraph, r: int, g: int, b: int) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _add_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.left_indent = Inches(0.25)


_STATUS_COLORS = {
    "pass": ("70AD47", "FFFFFF"),    # green, white text
    "warning": ("FFC000", "000000"), # amber, black text
    "fail": ("C00000", "FFFFFF"),    # red, white text
}


def _status_badge(cell, status: str) -> None:
    bg, fg = _STATUS_COLORS.get(status, ("CCCCCC", "000000"))
    _add_shading(cell, bg)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(fg))
            run.bold = True


def build_governance_docx(data: dict) -> bytes:
    """Build a governance report .docx from GovernanceAgent JSON output."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover ────────────────────────────────────────────────────────────────
    title = doc.add_heading("Governance & Quality Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, 31, 83, 141)
    doc.add_paragraph()

    # Overall score banner
    overall_status = data.get("status", "warning")
    overall_score = data.get("overall_score", 0)
    score_bg, score_fg = _STATUS_COLORS.get(overall_status, ("CCCCCC", "000000"))

    score_table = doc.add_table(rows=1, cols=2)
    score_table.style = "Table Grid"
    score_label = score_table.cell(0, 0)
    score_value = score_table.cell(0, 1)

    lrun = score_label.paragraphs[0].add_run("Overall Score / Status")
    lrun.bold = True
    _add_shading(score_label, "D6E4F0")

    vrun = score_value.paragraphs[0].add_run(
        f"{overall_score} / 10  —  {overall_status.upper()}"
    )
    vrun.bold = True
    _add_shading(score_value, score_bg)
    vrun.font.color.rgb = RGBColor(*bytes.fromhex(score_fg))

    doc.add_paragraph()

    # Summary
    h = doc.add_heading("Executive Summary", level=1)
    _set_heading_color(h, 31, 83, 141)
    doc.add_paragraph(data.get("summary", ""))

    # ── Checks Table ─────────────────────────────────────────────────────────
    h = doc.add_heading("Quality Checks", level=1)
    _set_heading_color(h, 31, 83, 141)

    checks = data.get("checks", [])
    if checks:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(["Check", "Status", "Score", "Findings"]):
            run = hdr_cells[i].paragraphs[0].add_run(header)
            run.bold = True
            _add_shading(hdr_cells[i], "1F538D")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for check in checks:
            row = table.add_row().cells
            row[0].paragraphs[0].add_run(check.get("check_name", ""))
            status_val = check.get("status", "warning")
            status_run = row[1].paragraphs[0].add_run(status_val.upper())
            status_run.bold = True
            _status_badge(row[1], status_val)
            row[2].paragraphs[0].add_run(str(check.get("score", "")))
            row[3].paragraphs[0].add_run(check.get("findings", ""))

            # Recommendations as sub-row
            recs = check.get("recommendations", [])
            if recs:
                rec_row = table.add_row().cells
                rec_row[0].merge(rec_row[3])
                rec_text = "  Recommendations: " + " | ".join(recs)
                rec_run = rec_row[0].paragraphs[0].add_run(rec_text)
                rec_run.italic = True
                rec_run.font.size = Pt(9)
                _add_shading(rec_row[0], "F5F5F5")

    doc.add_paragraph()

    # ── Missing Information ──────────────────────────────────────────────────
    missing = data.get("missing_information", [])
    if missing:
        h = doc.add_heading("Missing Information", level=1)
        _set_heading_color(h, 31, 83, 141)
        _bullet_list(doc, missing)

    # ── Contradictions ───────────────────────────────────────────────────────
    contradictions = data.get("contradictions", [])
    if contradictions:
        h = doc.add_heading("Contradictions Identified", level=1)
        _set_heading_color(h, 192, 0, 0)
        _bullet_list(doc, contradictions)

    # ── Risk Flags ───────────────────────────────────────────────────────────
    risk_flags = data.get("risk_flags", [])
    if risk_flags:
        h = doc.add_heading("Risk Flags", level=1)
        _set_heading_color(h, 192, 0, 0)
        _bullet_list(doc, risk_flags)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
