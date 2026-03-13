"""Word (.docx) formatter for the Persona & Research Agent output."""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


def _set_heading_color(paragraph, r: int, g: int, b: int) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _add_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.left_indent = Inches(0.25)


def _persona_card(doc: Document, persona: dict, index: int) -> None:
    """Render a single persona as a styled info card."""
    # Persona heading
    h = doc.add_heading(f"Persona {index}: {persona.get('name', 'Unknown')}", level=2)
    _set_heading_color(h, 46, 117, 182)

    # Metadata table
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    role_label = table.cell(0, 0)
    role_value = table.cell(0, 1)
    ctx_label = table.cell(1, 0)
    ctx_value = table.cell(1, 1)

    role_label.paragraphs[0].add_run("Role").bold = True
    _add_shading(role_label, "D6E4F0")
    role_value.paragraphs[0].add_run(persona.get("role", "—"))

    ctx_label.paragraphs[0].add_run("Context").bold = True
    _add_shading(ctx_label, "D6E4F0")
    ctx_value.paragraphs[0].add_run(persona.get("context", "—"))

    doc.add_paragraph()

    # Goals
    goals_h = doc.add_heading("Goals", level=3)
    _set_heading_color(goals_h, 70, 130, 180)
    _bullet_list(doc, persona.get("goals", []))

    # Pain Points
    pain_h = doc.add_heading("Pain Points", level=3)
    _set_heading_color(pain_h, 192, 0, 0)
    _bullet_list(doc, persona.get("pain_points", []))

    doc.add_paragraph()


def _interview_questions_table(doc: Document, questions: list[dict]) -> None:
    """Render all interview questions as a formatted table."""
    if not questions:
        doc.add_paragraph("No interview questions generated.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header
    hdr_cells = table.rows[0].cells
    headers = ["#", "Question", "Category", "Target Persona"]
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        _add_shading(hdr_cells[i], "1F538D")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Category colour map
    cat_colors = {
        "discovery": "4CAF50",
        "validation": "2196F3",
        "prioritization": "FF9800",
    }

    for idx, q in enumerate(questions, 1):
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(str(idx))
        row[1].paragraphs[0].add_run(q.get("question", ""))
        cat_run = row[2].paragraphs[0].add_run(q.get("category", ""))
        cat_run.bold = True
        cat_color = cat_colors.get(q.get("category", ""), "CCCCCC")
        _add_shading(row[2], cat_color)
        row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        row[3].paragraphs[0].add_run(q.get("target_persona") or "All")

        # Rationale as a sub-row note
        rationale = q.get("rationale", "")
        if rationale:
            note_row = table.add_row().cells
            note_row[0].merge(note_row[3])
            merged_para = note_row[0].paragraphs[0]
            merged_run = merged_para.add_run(f"  Rationale: {rationale}")
            merged_run.italic = True
            merged_run.font.size = Pt(9)
            _add_shading(note_row[0], "F5F5F5")


def build_personas_docx(personas: list[dict], interview_questions: list[dict]) -> bytes:
    """Build a persona research .docx from PersonaResearchAgent JSON output."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("User Personas & Research", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, 31, 83, 141)
    doc.add_paragraph()

    # ── Personas ─────────────────────────────────────────────────────────────
    h = doc.add_heading("User Personas", level=1)
    _set_heading_color(h, 31, 83, 141)

    if personas:
        for i, persona in enumerate(personas, 1):
            _persona_card(doc, persona, i)
    else:
        doc.add_paragraph("No personas generated.")

    # ── Interview Questions ───────────────────────────────────────────────────
    doc.add_page_break()
    h = doc.add_heading("Interview Questions", level=1)
    _set_heading_color(h, 31, 83, 141)
    p = doc.add_paragraph(
        f"Total questions: {len(interview_questions)}  |  "
        f"Personas covered: {len(personas)}"
    )
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    _interview_questions_table(doc, interview_questions)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
