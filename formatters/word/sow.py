"""Word (.docx) formatter for the SOW Generation Agent output."""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


def _set_heading_color(paragraph, r: int, g: int, b: int) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _add_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.left_indent = Inches(0.25)


def _info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """2-column label/value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        label_cell = table.cell(i, 0)
        value_cell = table.cell(i, 1)

        label_cell.width = Inches(2.0)
        label_run = label_cell.paragraphs[0].add_run(label)
        label_run.bold = True
        _add_shading(label_cell, "D6E4F0")

        value_cell.paragraphs[0].add_run(value or "—")


def build_sow_docx(data: dict) -> bytes:
    """Build a professional SOW .docx from SOWAgent JSON output."""
    doc = Document()

    # Document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover Title ──────────────────────────────────────────────────────────
    title = doc.add_heading(data.get("project_title", "Statement of Work"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, 31, 83, 141)

    subtitle = doc.add_paragraph("Statement of Work")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(46, 117, 182)

    doc.add_paragraph()  # spacer

    # Metadata table
    _info_table(doc, [
        ("Estimated Effort", data.get("estimated_effort", "—")),
    ])
    doc.add_paragraph()

    # ── Executive Summary ────────────────────────────────────────────────────
    h = doc.add_heading("1. Executive Summary", level=1)
    _set_heading_color(h, 31, 83, 141)
    doc.add_paragraph(data.get("executive_summary", ""))

    # ── Scope of Work ────────────────────────────────────────────────────────
    scope = data.get("scope", {})
    h = doc.add_heading("2. Scope of Work", level=1)
    _set_heading_color(h, 31, 83, 141)
    doc.add_paragraph(scope.get("content", ""))

    # ── Deliverables ─────────────────────────────────────────────────────────
    h = doc.add_heading("3. Deliverables", level=1)
    _set_heading_color(h, 31, 83, 141)
    _bullet_list(doc, data.get("deliverables", []))

    # ── Timeline ─────────────────────────────────────────────────────────────
    timeline = data.get("timeline", {})
    h = doc.add_heading("4. Project Timeline", level=1)
    _set_heading_color(h, 31, 83, 141)
    doc.add_paragraph(timeline.get("content", ""))

    # ── Assumptions ──────────────────────────────────────────────────────────
    h = doc.add_heading("5. Assumptions", level=1)
    _set_heading_color(h, 31, 83, 141)
    _bullet_list(doc, data.get("assumptions", []))

    # ── Constraints ──────────────────────────────────────────────────────────
    h = doc.add_heading("6. Constraints", level=1)
    _set_heading_color(h, 31, 83, 141)
    _bullet_list(doc, data.get("constraints", []))

    # ── Acceptance Criteria ──────────────────────────────────────────────────
    h = doc.add_heading("7. Acceptance Criteria", level=1)
    _set_heading_color(h, 31, 83, 141)
    _bullet_list(doc, data.get("acceptance_criteria", []))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
