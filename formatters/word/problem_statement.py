"""Word (.docx) formatter for the Problem Statement Agent output."""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


def _set_heading_color(paragraph, r: int, g: int, b: int) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _add_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.left_indent = Inches(0.25)


def _labeled_section(doc: Document, label: str, content: str, level: int = 2) -> None:
    h = doc.add_heading(label, level=level)
    _set_heading_color(h, 46, 117, 182)
    if content:
        doc.add_paragraph(content)


def build_problem_statement_docx(data: dict) -> bytes:
    """Build a structured problem statement .docx from ProblemStatementAgent JSON output."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading(data.get("problem_title", "Problem Statement"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, 31, 83, 141)

    subtitle = doc.add_paragraph("Problem Statement Document")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    doc.add_paragraph()

    # ── Problem Statement ────────────────────────────────────────────────────
    h = doc.add_heading("1. Problem Statement", level=1)
    _set_heading_color(h, 31, 83, 141)
    doc.add_paragraph(data.get("problem_statement", ""))

    # ── State Analysis ───────────────────────────────────────────────────────
    h = doc.add_heading("2. State Analysis", level=1)
    _set_heading_color(h, 31, 83, 141)

    _labeled_section(doc, "2.1 Current State", data.get("current_state", ""))
    _labeled_section(doc, "2.2 Desired State", data.get("desired_state", ""))
    _labeled_section(doc, "2.3 Gap Analysis", data.get("gap_analysis", ""))

    # ── Stakeholders ─────────────────────────────────────────────────────────
    h = doc.add_heading("3. Stakeholders Affected", level=1)
    _set_heading_color(h, 31, 83, 141)
    _bullet_list(doc, data.get("stakeholders_affected", []))

    # ── Business Impact ──────────────────────────────────────────────────────
    h = doc.add_heading("4. Business Impact", level=1)
    _set_heading_color(h, 31, 83, 141)
    doc.add_paragraph(data.get("business_impact", ""))

    # ── Constraints ──────────────────────────────────────────────────────────
    h = doc.add_heading("5. Constraints", level=1)
    _set_heading_color(h, 31, 83, 141)
    _bullet_list(doc, data.get("constraints", []))

    # ── Success Vision ───────────────────────────────────────────────────────
    h = doc.add_heading("6. Success Vision", level=1)
    _set_heading_color(h, 31, 83, 141)
    p = doc.add_paragraph(data.get("success_vision", ""))
    p.runs[0].italic = True if p.runs else None

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
